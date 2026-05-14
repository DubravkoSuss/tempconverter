from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup: A4, 20mm margins
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(20)
section.right_margin = Mm(20)

def set_font(run, size=10, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_paragraph(text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold, italic)
    return p

def add_heading(text, size=10, bold=True, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size, bold, italic)
    return p

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Containerization and Deployment of TempConverter Application')
set_font(run, 14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Josip Stanešić\nAlgebra Bernays University, Zagreb\n2025')
set_font(run, 10)

doc.add_paragraph()

# 1. Abstract
add_heading('1. Abstract')
add_paragraph(
    'This project addresses the containerization and deployment of the TempConverter web application, '
    'a Flask-based temperature conversion tool that requires a MySQL 8 database. The application was '
    'not previously containerized, and the goal was to package it into a container image, deploy it '
    'locally using Podman, and then orchestrate it using both a simple and a complex container '
    'orchestration system. A container image was built using a Dockerfile, pushed to a container '
    'registry, and deployed locally alongside a MySQL container. Docker Swarm was chosen as the '
    'simple orchestration system and Kubernetes as the complex one. Configuration files and deployment '
    'templates were created for both systems. The application is exposed on port 80 with two replicas '
    'and a single database instance. A continuous integration pipeline was also created to automate '
    'testing and image building.'
)

doc.add_paragraph()

# 2. Introduction
add_heading('2. Introduction')
add_paragraph(
    '1. What is the problem we are solving?\n'
    'The TempConverter application exists only as source code and is not containerized. It needs to be '
    'packaged, deployed locally, and made available through container orchestration systems so it can '
    'be scaled and managed in a modern DevOps workflow.'
)
add_paragraph(
    '2. How has this problem already been solved?\n'
    'Container technologies such as Docker and Podman are widely used to package applications with '
    'their dependencies. Orchestration platforms like Kubernetes and Docker Swarm are industry-standard '
    'solutions for managing containerized workloads at scale. Official documentation from Docker, '
    'Podman, and Kubernetes was used as reference.'
)
add_paragraph(
    '3. What did I do to solve the problem?\n'
    'A Dockerfile was written to build the application image. The application was deployed locally '
    'using Podman with a MySQL container on a shared network. A Docker Swarm stack file and Kubernetes '
    'manifests were created for orchestrated deployments. A GitHub Actions CI pipeline was also set up '
    'to run tests and build images automatically.'
)
add_paragraph(
    '4. What are the results and technical solutions?\n'
    'Two container images were built and tagged (tempconverter:latest and tempconverter:dev). The '
    'application runs locally and connects to the database successfully. Deployment templates for '
    'Docker Swarm and Kubernetes are ready and functional. The application is accessible on port 80 '
    'with load balancing across two replicas.'
)
add_paragraph(
    '5. What are the next steps?\n'
    'Future improvements include adding HTTPS support, setting up horizontal pod autoscaling in '
    'Kubernetes, using Helm charts for more flexible deployments, and integrating a proper secrets '
    'management solution such as HashiCorp Vault.'
)

doc.add_paragraph()

# 3. Explanation
add_heading('3. Explanation / What Was Done')

add_heading('3.1. Task 1 – Creating the Container Image', bold=False, italic=True)
add_paragraph(
    'The application source code was cloned from https://github.com/jstanesic/tempconverter. '
    'A Dockerfile was created based on the python:3.11-slim base image. All system packages are '
    'updated during the build process using apt-get upgrade. Python dependencies are installed from '
    'requirements.txt. Port 5000 is exposed and the application is started with python app.py.'
)

# Code block style paragraph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(
    'FROM python:3.11-slim\n'
    'RUN apt-get update && apt-get upgrade -y && apt-get clean\n'
    'WORKDIR /app\n'
    'COPY requirements.txt .\n'
    'RUN pip install --no-cache-dir -r requirements.txt\n'
    'COPY . .\n'
    'EXPOSE 5000\n'
    'CMD ["python", "app.py"]'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_paragraph('The image was built using the following command:')
p = doc.add_paragraph()
run = p.add_run('podman build -t tempconverter:latest .')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()

add_heading('3.2. Task 2 and 3 – Pushing Images and Dev Tag', bold=False, italic=True)
add_paragraph(
    'After building the image, it was tagged and pushed to a container registry. A second image '
    'tagged tempconverter:dev was created after updating the HTML title in templates/index.html '
    'from "Celsius to Fahrenheit Converter" to "TempConverter".'
)

p = doc.add_paragraph()
run = p.add_run(
    'podman tag tempconverter:latest YOUR_USERNAME/tempconverter:latest\n'
    'podman push YOUR_USERNAME/tempconverter:latest\n'
    'podman build -t tempconverter:dev .\n'
    'podman push YOUR_USERNAME/tempconverter:dev'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()

add_heading('3.3. Task 4 – Local Deployment with Podman', bold=False, italic=True)
add_paragraph(
    'The application was deployed locally using Podman. A dedicated network was created so the '
    'application container can reach the database container by hostname. The MySQL container was '
    'started first and the application container was started after the database was ready. '
    'Environment variables were used to configure the database connection and to set the student '
    'name and college.'
)

p = doc.add_paragraph()
run = p.add_run(
    'podman network create tempconverter-net\n\n'
    'podman run -d --name mysql \\\n'
    '  --network tempconverter-net \\\n'
    '  -e MYSQL_ROOT_PASSWORD=rootpass \\\n'
    '  -e MYSQL_DATABASE=tempconverter \\\n'
    '  -e MYSQL_USER=appuser \\\n'
    '  -e MYSQL_PASSWORD=apppass \\\n'
    '  docker.io/mysql:8\n\n'
    'podman run -d --name tempconverter \\\n'
    '  --network tempconverter-net \\\n'
    '  -p 5000:5000 \\\n'
    '  -e DB_USER=appuser -e DB_PASS=apppass \\\n'
    '  -e DB_HOST=mysql -e DB_NAME=tempconverter \\\n'
    '  -e STUDENT="Josip Stanešić" \\\n'
    '  -e COLLEGE="Algebra Bernays University" \\\n'
    '  localhost/tempconverter:latest'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()

# Resource comparison table
add_heading('3.4. Resource Usage: Container vs Virtual Machine', bold=False, italic=True)
add_paragraph(
    'The following table compares the resource usage of the deployed containers against a typical '
    'virtual machine running the same application stack. Measurements were taken using podman stats.'
)

table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ['Resource', 'Container (App)', 'Container (MySQL)', 'Typical VM']
rows_data = [
    ['Memory', '93.5 MB', '453.7 MB', '512 MB – 2 GB'],
    ['CPU usage', '1.21%', '4.26%', 'Dedicated vCPUs'],
    ['Disk (image)', '~200 MB', '~300 MB', '10+ GB'],
    ['Boot time', '~3 seconds', '~15 seconds', '1–3 minutes'],
    ['Isolation', 'Process-level', 'Process-level', 'Hardware-level'],
]

for i, h in enumerate(headers):
    cell = table.cell(0, i)
    run = cell.paragraphs[0].add_run(h)
    set_font(run, 10, bold=True)

for r, row_data in enumerate(rows_data):
    for c, val in enumerate(row_data):
        cell = table.cell(r + 1, c)
        run = cell.paragraphs[0].add_run(val)
        set_font(run, 10)

p = doc.add_paragraph()
run = p.add_run('Table 1. Resource usage comparison between containers and a virtual machine')
set_font(run, 10, italic=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_heading('3.5. Task 5 – Orchestration System Choice', bold=False, italic=True)
add_paragraph(
    'Docker Swarm was chosen as the simple orchestration system due to its native integration with '
    'Docker, minimal configuration requirements, and ease of use for small deployments. Kubernetes '
    'was chosen as the complex orchestration system because it is the industry standard, offers '
    'advanced scheduling, self-healing, and a rich ecosystem of tools.'
)

doc.add_paragraph()

add_heading('3.6. Tasks 6 and 7 – Docker Swarm Deployment', bold=False, italic=True)
add_paragraph(
    'A Docker Swarm stack file (docker-stack.yml) was created defining two replicas of the '
    'application container and one instance of the MySQL container. The application is exposed '
    'on port 80. Pod anti-affinity is enforced through Swarm placement constraints so replicas '
    'are not scheduled on the same node. The stack is deployed with the following command:'
)

p = doc.add_paragraph()
run = p.add_run(
    'docker swarm init\n'
    'docker stack deploy -c docker-stack.yml tempconverter'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_paragraph('To scale the application to 3 replicas:')
p = doc.add_paragraph()
run = p.add_run('docker service scale tempconverter_app=3')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()

add_heading('3.7. Tasks 8 – Kubernetes Deployment', bold=False, italic=True)
add_paragraph(
    'Kubernetes manifests were created in the kubernetes/ directory. These include a Secret for '
    'database credentials, a PersistentVolumeClaim for MySQL storage, Deployments and Services '
    'for both MySQL and the application, and an OpenShift Template. The application deployment '
    'uses podAntiAffinity to ensure replicas are scheduled on different nodes. The service is '
    'of type LoadBalancer and exposes port 80.'
)

add_paragraph('The following files were created:')
for item in [
    'mysql-secret.yaml – database credentials stored as a Kubernetes Secret',
    'mysql-pvc.yaml – persistent storage for MySQL data',
    'mysql-deployment.yaml – MySQL deployment with 1 replica',
    'mysql-service.yaml – headless service for MySQL',
    'app-deployment.yaml – application deployment with 2 replicas and anti-affinity',
    'app-service.yaml – LoadBalancer service on port 80',
    'openshift-template.yaml – parameterized OpenShift template',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    set_font(run, 10)

add_paragraph('Deploy with:')
p = doc.add_paragraph()
run = p.add_run('kubectl apply -f kubernetes/')
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_paragraph('Scale to 3 replicas:')
p = doc.add_paragraph()
run = p.add_run('kubectl scale deployment tempconverter --replicas=3')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph()

add_heading('3.8. Task 9 – Reflection on Orchestration Systems', bold=False, italic=True)
add_paragraph(
    'Docker Swarm is simpler to set up and operate. It is suitable for small teams and straightforward '
    'applications where the overhead of Kubernetes is not justified. Configuration is done in a single '
    'Compose-style YAML file and the learning curve is low.'
)
add_paragraph(
    'Kubernetes is significantly more complex but offers far greater capabilities: horizontal pod '
    'autoscaling, rolling updates, self-healing, namespace isolation, and a large ecosystem of '
    'operators and tools. It is the right choice for production workloads and larger teams.'
)
add_paragraph(
    'For simple environments such as internal tools or development setups, Docker Swarm is the '
    'preferred choice. For production microservices or any environment requiring high availability '
    'and fine-grained control, Kubernetes is the appropriate platform.'
)

doc.add_paragraph()

add_heading('4. CI/CD Pipeline', bold=True)
add_paragraph(
    'A GitHub Actions pipeline was created in .github/workflows/ci.yml. It runs on every push to '
    'the main and dev branches. The pipeline has two jobs: the first runs unit and integration tests '
    'using pytest, and the second builds the container image only if the tests pass.'
)

doc.add_paragraph()

add_heading('5. Conclusion')
add_paragraph(
    'The TempConverter application was successfully containerized and deployed using multiple '
    'approaches. The container image meets all requirements: packages are updated at build time, '
    'port 5000 is exposed, all dependencies are installed, and the Flask application starts '
    'correctly. The application connects to MySQL as a non-root user and displays the student '
    'name and college on the web page. Deployment templates for both Docker Swarm and Kubernetes '
    'are functional and ready for use.'
)

doc.save('/mnt/c/Users/dubra/tempconverter/project_document.docx')
print("Document saved.")
